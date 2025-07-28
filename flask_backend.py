#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Flask API Backend for Company News Analyzer Web App
Enhanced for bulk analysis of multiple companies
"""

from flask import Flask, request, jsonify, send_file
from flask_cors import CORS
import asyncio
import json
import os
import threading
import time
from datetime import datetime
import traceback
import tempfile

# Import your existing analysis code
try:
    from streamlined_analyzer import CompanyAnalyzer, NEWSPAPERS
    print("✅ Successfully imported analysis code")
except ImportError as e:
    print(f"❌ Error importing analysis code: {e}")
    print("Make sure 'streamlined_analyzer.py' exists in the same folder")
    exit(1)

app = Flask(__name__)
CORS(app)  # Enable CORS for React frontend

# Global variable to store analysis progress
analysis_sessions = {}

# Pre-defined companies list (matches frontend)
PREDEFINED_COMPANIES = [
    'Gulf Bank',
    'National Bank of Kuwait', 
    'Commercial Bank of Kuwait',
    'Burgan Bank',
    'Kuwait Finance House',
    'Kuwait International Bank',
    'Al Ahli Bank of Kuwait',
    'Warba Bank'
]

class AnalysisStatus:
    def __init__(self):
        self.current_step = ""
        self.completed_newspapers = []
        self.total_newspapers = 0
        self.is_running = False
        self.results = None
        self.error = None
        self.progress_percentage = 0
        self.analysis_type = "detailed"
        self.report_path = None
        self.company_name = ""

@app.route('/api/health', methods=['GET'])
def health_check():
    """Health check endpoint"""
    return jsonify({
        "status": "healthy", 
        "timestamp": datetime.now().isoformat(),
        "message": "Company News Analyzer API is running",
        "predefined_companies": len(PREDEFINED_COMPANIES)
    })

@app.route('/api/companies', methods=['GET'])
def get_companies():
    """Get list of pre-defined companies"""
    try:
        companies_list = []
        for i, name in enumerate(PREDEFINED_COMPANIES):
            # Create safe ID for frontend
            safe_id = name.lower().replace(' ', '_').replace('-', '_')
            companies_list.append({
                "id": safe_id,
                "name": name,
                "selected": True  # All selected by default
            })
        
        return jsonify({
            "companies": companies_list,
            "total": len(companies_list)
        })
    except Exception as e:
        return jsonify({"error": f"Failed to load companies: {str(e)}"}), 500

@app.route('/api/newspapers', methods=['GET'])
def get_newspapers():
    """Get list of available newspapers"""
    try:
        newspapers_list = []
        for name, config in NEWSPAPERS.items():
            # Create safe ID for frontend
            safe_id = name.lower().replace(' ', '_').replace('-', '_')
            newspapers_list.append({
                "id": safe_id,
                "name": name,
                "language": config.get("language", "Arabic"),
                "url": config.get("url", "")
            })
        
        return jsonify({
            "newspapers": newspapers_list,
            "total": len(newspapers_list)
        })
    except Exception as e:
        return jsonify({"error": f"Failed to load newspapers: {str(e)}"}), 500

@app.route('/api/analyze', methods=['POST'])
def start_analysis():
    """Start company analysis (now supports single company from bulk request)"""
    try:
        data = request.get_json()
        
        if not data:
            return jsonify({"error": "No data provided"}), 400
            
        company_name = data.get('companyName', '').strip()
        selected_newspaper_ids = data.get('selectedNewspapers', [])
        analysis_type = data.get('analysisType', 'detailed')
        
        # Validation
        if not company_name:
            return jsonify({"error": "Company name is required"}), 400
        
        if not selected_newspaper_ids:
            return jsonify({"error": "At least one newspaper must be selected"}), 400
        
        if analysis_type not in ['short', 'detailed']:
            return jsonify({"error": "Invalid analysis type"}), 400
        
        # Validate company is in predefined list
        if company_name not in PREDEFINED_COMPANIES:
            return jsonify({"error": f"Company '{company_name}' is not in the predefined list"}), 400
        
        print(f"📝 Starting {analysis_type} analysis for: {company_name}")
        print(f"📰 Selected newspapers: {selected_newspaper_ids}")
        
        # Map newspaper IDs back to newspaper names
        id_to_name_mapping = {
            name.lower().replace(' ', '_').replace('-', '_'): name 
            for name in NEWSPAPERS.keys()
        }
        
        selected_newspapers = {}
        for newspaper_id in selected_newspaper_ids:
            newspaper_name = id_to_name_mapping.get(newspaper_id)
            if newspaper_name and newspaper_name in NEWSPAPERS:
                selected_newspapers[newspaper_name] = NEWSPAPERS[newspaper_name]
        
        if not selected_newspapers:
            return jsonify({"error": "No valid newspapers selected"}), 400
        
        print(f"✅ Mapped to newspapers: {list(selected_newspapers.keys())}")
        
        # Create analysis session
        session_id = f"{company_name.replace(' ', '_')}_{int(time.time())}"
        analysis_sessions[session_id] = AnalysisStatus()
        analysis_sessions[session_id].total_newspapers = len(selected_newspapers)
        analysis_sessions[session_id].is_running = True
        analysis_sessions[session_id].current_step = "Initializing analysis..."
        analysis_sessions[session_id].analysis_type = analysis_type
        analysis_sessions[session_id].company_name = company_name
        
        print(f"🚀 Created session: {session_id}")
        
        # Start analysis in background thread
        def run_analysis():
            try:
                print(f"🔄 Starting background analysis for session: {session_id}")
                loop = asyncio.new_event_loop()
                asyncio.set_event_loop(loop)
                result = loop.run_until_complete(
                    analyze_company_async(session_id, company_name, selected_newspapers, analysis_type)
                )
                analysis_sessions[session_id].results = result
                analysis_sessions[session_id].is_running = False
                analysis_sessions[session_id].current_step = "Analysis completed successfully!"
                analysis_sessions[session_id].progress_percentage = 100
                print(f"✅ Analysis completed for session: {session_id}")
                loop.close()
            except Exception as e:
                error_msg = str(e)
                analysis_sessions[session_id].error = error_msg
                analysis_sessions[session_id].is_running = False
                analysis_sessions[session_id].current_step = f"Error: {error_msg}"
                print(f"❌ Analysis failed for session {session_id}: {error_msg}")
                traceback.print_exc()
        
        # Start analysis thread
        analysis_thread = threading.Thread(target=run_analysis)
        analysis_thread.daemon = True
        analysis_thread.start()
        
        return jsonify({
            "sessionId": session_id,
            "message": f"{analysis_type.title()} analysis started successfully",
            "companyName": company_name,
            "selectedNewspapers": list(selected_newspapers.keys()),
            "totalNewspapers": len(selected_newspapers),
            "analysisType": analysis_type
        })
        
    except Exception as e:
        print(f"❌ Error in start_analysis: {e}")
        traceback.print_exc()
        return jsonify({"error": f"Failed to start analysis: {str(e)}"}), 500

@app.route('/api/analysis/<session_id>/status', methods=['GET'])
def get_analysis_status(session_id):
    """Get analysis progress status"""
    if session_id not in analysis_sessions:
        return jsonify({"error": "Session not found"}), 404
    
    session = analysis_sessions[session_id]
    
    return jsonify({
        "isRunning": session.is_running,
        "currentStep": session.current_step,
        "completedNewspapers": session.completed_newspapers,
        "totalNewspapers": session.total_newspapers,
        "progressPercentage": session.progress_percentage,
        "hasResults": session.results is not None,
        "hasError": session.error is not None,
        "error": session.error,
        "analysisType": session.analysis_type,
        "companyName": session.company_name
    })

@app.route('/api/analysis/<session_id>/results', methods=['GET'])
def get_analysis_results(session_id):
    """Get analysis results"""
    if session_id not in analysis_sessions:
        return jsonify({"error": "Session not found"}), 404
    
    session = analysis_sessions[session_id]
    
    if session.error:
        return jsonify({"error": session.error}), 500
    
    if not session.results:
        return jsonify({"error": "Results not ready yet"}), 202
    
    return jsonify(session.results)

@app.route('/api/analysis/<session_id>/stop', methods=['POST'])
def stop_analysis(session_id):
    """Stop a running analysis"""
    if session_id not in analysis_sessions:
        return jsonify({"error": "Session not found"}), 404
    
    session = analysis_sessions[session_id]
    
    try:
        # Mark the session as stopped
        session.is_running = False
        session.current_step = "Analysis stopped by user"
        session.error = "Stopped by user"
        
        print(f"🛑 Analysis stopped by user for session: {session_id} ({session.company_name})")
        
        return jsonify({
            "message": f"Analysis stopped for {session.company_name}",
            "sessionId": session_id,
            "status": "stopped"
        })
        
    except Exception as e:
        print(f"❌ Error stopping analysis {session_id}: {e}")
        return jsonify({"error": f"Failed to stop analysis: {str(e)}"}), 500
    """Download analysis report"""
    if session_id not in analysis_sessions:
        return jsonify({"error": "Session not found"}), 404
    
    session = analysis_sessions[session_id]
    
    if not session.results:
        return jsonify({"error": "Results not ready"}), 202
    
    try:
        if hasattr(session, 'report_path') and session.report_path:
            report_path = session.report_path
            
            if os.path.exists(report_path):
                # Extract filename from path
                original_filename = os.path.basename(report_path)
                
                print(f"📥 Serving Word report: {original_filename}")
                
                return send_file(
                    report_path,
                    as_attachment=True,
                    download_name=original_filename,
                    mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
                )
            else:
                print(f"❌ Word report not found at: {report_path}")
                return jsonify({"error": "Word report file not found"}), 404
        else:
            print(f"❌ No Word report found for session {session_id}")
            return jsonify({"error": "No report available for download"}), 404
            
    except Exception as e:
        print(f"❌ Download error: {e}")
        return jsonify({"error": f"Failed to create download: {str(e)}"}), 500

@app.route('/api/headlines-report', methods=['POST'])
def generate_headlines_report():
    """Generate consolidated headlines report for multiple completed analyses"""
    try:
        data = request.get_json()
        session_ids = data.get('sessionIds', [])
        
        if not session_ids:
            return jsonify({"error": "No session IDs provided"}), 400
        
        # Collect results from all sessions
        bank_headlines = {}
        analysis_type = "detailed"  # Default
        
        for session_id in session_ids:
            if session_id in analysis_sessions:
                session = analysis_sessions[session_id]
                if session.results and hasattr(session, 'company_name'):
                    bank_headlines[session.company_name] = {
                        'results': session.results,
                        'session': session
                    }
                    analysis_type = session.analysis_type
        
        if not bank_headlines:
            return jsonify({"error": "No valid completed analyses found"}), 400
        
        # Generate headlines report
        print(f"📰 Generating headlines report for {len(bank_headlines)} banks")
        
        # Run the async function in a new event loop
        def run_headlines_generation():
            loop = asyncio.new_event_loop()
            asyncio.set_event_loop(loop)
            try:
                report_path = loop.run_until_complete(
                    generate_consolidated_headlines_report(bank_headlines, analysis_type)
                )
                return report_path
            finally:
                loop.close()
        
        report_path = run_headlines_generation()
        
        if report_path and os.path.exists(report_path):
            filename = os.path.basename(report_path)
            
            return send_file(
                report_path,
                as_attachment=True,
                download_name=filename,
                mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            )
        else:
            return jsonify({"error": "Failed to generate headlines report"}), 500
            
    except Exception as e:
        print(f"❌ Headlines report error: {e}")
        traceback.print_exc()
        return jsonify({"error": f"Failed to generate headlines report: {str(e)}"}), 500

async def analyze_company_async(session_id, company_name, selected_newspapers, analysis_type):
    """Async wrapper for company analysis"""
    try:
        session = analysis_sessions[session_id]
        print(f"🔄 Starting async {analysis_type} analysis for: {company_name}")
        
        # Initialize analyzer
        session.current_step = "Initializing analyzer..."
        session.progress_percentage = 5
        analyzer = CompanyAnalyzer()
        print("✅ Analyzer initialized")
        
        # Generate variations
        session.current_step = f"Generating AI variations for: {company_name}"
        session.progress_percentage = 10
        print(f"🤖 Generating variations for: {company_name}")
        variations = await analyzer.generate_company_variations(company_name)
        print(f"✅ Generated variations: {len(variations.get('english_variations', []))} English, {len(variations.get('arabic_variations', []))} Arabic")
        
        session.current_step = "Starting newspaper analysis..."
        session.progress_percentage = 15
        
        all_results = []
        total_newspapers = len(selected_newspapers)
        
        # Analyze each newspaper
        for i, (newspaper_name, config) in enumerate(selected_newspapers.items()):
            # Check if analysis was stopped
            if not session.is_running:
                print(f"🛑 Analysis stopped for {company_name} during {newspaper_name}")
                break
                
            try:
                session.current_step = f"Analyzing {newspaper_name}..."
                session.progress_percentage = 15 + (i / total_newspapers) * 70
                print(f"📰 Analyzing {newspaper_name} ({i+1}/{total_newspapers})")
                
                result = await analyzer.analyze_newspaper(newspaper_name, config, company_name, variations)
                all_results.append(result)
                session.completed_newspapers.append(newspaper_name)
                
                if result["success"] and result["relevant_pages"]:
                    print(f"✅ {newspaper_name}: Found content on {len(result['relevant_pages'])} pages")
                elif result["success"]:
                    print(f"✅ {newspaper_name}: No content found")
                else:
                    print(f"❌ {newspaper_name}: Failed - {result.get('error', 'Unknown error')}")
                
                # Brief pause for UI updates and stop check
                await asyncio.sleep(0.1)
                
            except Exception as e:
                print(f"❌ Error analyzing {newspaper_name}: {e}")
                all_results.append({
                    "newspaper": newspaper_name,
                    "success": False,
                    "error": str(e),
                    "relevant_pages": [],
                    "analysis": "",
                    "language": config.get("language", "N/A")
                })
        
        # Check if analysis was stopped before generating summary
        if not session.is_running:
            print(f"🛑 Analysis stopped for {company_name} before generating summary")
            session.current_step = "Analysis stopped by user"
            raise Exception("Analysis stopped by user")
        
        # Generate executive summary
        session.current_step = "Generating executive summary..."
        session.progress_percentage = 90
        print("📝 Generating executive summary...")
        
        successful_analyses = [r for r in all_results if r["success"] and r["relevant_pages"]]
        
        if analysis_type == "short":
            executive_summary = await generate_short_summary(analyzer, company_name, successful_analyses, len(selected_newspapers))
        else:
            executive_summary = await generate_detailed_summary(analyzer, company_name, successful_analyses, len(selected_newspapers))
        
        # Generate Word report
        session.current_step = "Creating Word report..."
        session.progress_percentage = 95
        print("📝 Creating Word report...")
        
        report_path = await analyzer.create_report(all_results, company_name, analysis_type)
        print(f"✅ Report saved: {report_path}")
        
        # Clean up temporary files
        await analyzer.cleanup_files(all_results)
        print("🧹 Temporary files cleaned up")
        
        # Calculate statistics
        successful_count = sum(1 for r in all_results if r["success"])
        content_count = sum(1 for r in all_results if r["success"] and r["relevant_pages"])
        total_pages = sum(len(r["relevant_pages"]) for r in all_results if r["success"])
        
        # Format results for web app
        if analysis_type == "short":
            web_results = create_short_web_results(company_name, executive_summary, all_results, successful_count, content_count, total_pages, len(selected_newspapers))
        else:
            web_results = create_detailed_web_results(company_name, executive_summary, all_results, successful_count, content_count, total_pages, len(selected_newspapers), report_path)
        
        # Store report path in session
        if report_path:
            analysis_sessions[session_id].report_path = report_path
        
        session.current_step = f"{analysis_type.title()} analysis completed!"
        session.progress_percentage = 100
        
        print(f"🎉 {analysis_type.title()} analysis completed for {company_name}!")
        print(f"📊 Statistics: {successful_count}/{len(selected_newspapers)} successful, {content_count} with content, {total_pages} total pages")
        
        return web_results
        
    except Exception as e:
        session.error = str(e)
        session.current_step = f"Error: {str(e)}"
        print(f"❌ Analysis error: {e}")
        traceback.print_exc()
        raise

async def generate_short_summary(analyzer, company_name, successful_analyses, total_newspapers):
    """Generate short summary - just bullet points"""
    if not successful_analyses:
        return {
            "keyDevelopments": f"No {company_name} content found in the analyzed newspapers.",
            "sentiment": "Neutral",
            "insights": f"Analysis completed across {total_newspapers} newspapers with no relevant content detected."
        }
    
    try:
        # Create simple bullet point summary
        all_content = []
        for result in successful_analyses:
            # Extract main points from each newspaper's analysis
            analysis_text = result['analysis'][:500]  # Limit content
            all_content.append(f"{result['newspaper']}: {analysis_text}")
        
        summary_prompt = f"""
        Create a BRIEF executive summary for {company_name}'s media coverage.
        Keep it to ONE paragraph with bullet points of main news items.
        
        Content from newspapers:
        {chr(10).join(all_content)}
        
        Provide a single paragraph summary with bullet points of the main news items about {company_name}.
        Focus only on the key facts, no analysis or sentiment.
        """
        
        summary_result = await analyzer.call_llm_with_retry(summary_prompt)
        
        return {
            "keyDevelopments": summary_result.strip(),
            "sentiment": "Not analyzed in short mode",
            "insights": f"Found {company_name} coverage in {len(successful_analyses)} out of {total_newspapers} newspapers."
        }
        
    except Exception as e:
        print(f"⚠️ Could not generate short summary: {e}")
        return {
            "keyDevelopments": f"Found {company_name} content in {len(successful_analyses)} newspapers with {sum(len(r['relevant_pages']) for r in successful_analyses)} total relevant pages.",
            "sentiment": "Not analyzed",
            "insights": f"Short analysis completed successfully across {total_newspapers} newspapers."
        }

async def generate_detailed_summary(analyzer, company_name, successful_analyses, total_newspapers):
    """Generate detailed summary with insights and sentiment"""
    if not successful_analyses:
        return {
            "keyDevelopments": f"Analysis of {company_name} across {total_newspapers} newspapers completed. No relevant content found.",
            "sentiment": "Neutral",
            "insights": f"Comprehensive analysis shows no current {company_name} coverage in the analyzed publications."
        }
    
    try:
        all_content = [f"{r['newspaper']}: {r['analysis']}" for r in successful_analyses]
        exec_prompt = f"""
        Create comprehensive executive summary for {company_name}'s media coverage:
        {chr(10).join(all_content)}
        
        Provide:
        KEY DEVELOPMENTS: [Main stories about {company_name} in 2-3 sentences]
        SENTIMENT: [Positive/Negative/Neutral with brief explanation]
        INSIGHTS: [Key business implications and recommendations in 2-3 sentences]
        """
        
        exec_result = await analyzer.call_llm_with_retry(exec_prompt)
        
        # Parse AI response
        executive_summary = {
            "keyDevelopments": f"Media analysis shows {company_name} presence across {len(successful_analyses)} out of {total_newspapers} newspapers.",
            "sentiment": "Neutral",
            "insights": f"Found coverage in {len(successful_analyses)} publications with {sum(len(r['relevant_pages']) for r in successful_analyses)} total relevant pages."
        }
        
        if "KEY DEVELOPMENTS:" in exec_result and "SENTIMENT:" in exec_result:
            parts = exec_result.split("SENTIMENT:")
            if len(parts) >= 2:
                executive_summary["keyDevelopments"] = parts[0].replace("KEY DEVELOPMENTS:", "").strip()
                
                sentiment_part = parts[1].split("INSIGHTS:")
                if len(sentiment_part) >= 2:
                    executive_summary["sentiment"] = sentiment_part[0].strip()
                    executive_summary["insights"] = sentiment_part[1].strip()
        
        print("✅ AI executive summary generated")
        return executive_summary
        
    except Exception as e:
        print(f"⚠️ Could not generate detailed summary: {e}")
        return {
            "keyDevelopments": f"Analysis of {company_name} across {total_newspapers} newspapers completed. Found relevant content in {len(successful_analyses)} publications.",
            "sentiment": "Neutral",
            "insights": f"Media analysis shows {company_name} presence across {len(successful_analyses)} out of {total_newspapers} newspapers with {sum(len(r['relevant_pages']) for r in successful_analyses)} total relevant pages."
        }

def create_short_web_results(company_name, executive_summary, all_results, successful_count, content_count, total_pages, total_newspapers):
    """Create web results for short analysis"""
    return {
        "companyName": company_name,
        "analysisDate": datetime.now().strftime('%Y-%m-%d'),
        "analysisType": "short",
        "executiveSummary": executive_summary,
        "newspaperResults": [
            {
                "id": result["newspaper"].lower().replace(' ', '_').replace('-', '_'),
                "newspaper": result["newspaper"],
                "language": result.get("language", "N/A"),
                "success": result["success"],
                "pagesFound": len(result["relevant_pages"]) if result["success"] else 0,
                "status": (
                    "Content Found" if result["success"] and result["relevant_pages"] else
                    "No Content" if result["success"] else
                    "Failed"
                ),
                "summary": result.get("analysis", result.get("error", "No analysis available"))[:200] + "..." if len(result.get("analysis", "")) > 200 else result.get("analysis", result.get("error", "No analysis available"))
            }
            for result in all_results
        ],
        "statistics": {
            "totalNewspapers": total_newspapers,
            "successfulAnalyses": successful_count,
            "newspapersWithContent": content_count,
            "totalRelevantPages": total_pages
        }
    }

async def generate_consolidated_headlines_report(bank_headlines, analysis_type):
    """Generate consolidated headlines report for all banks"""
    try:
        from docx import Document
        from docx.shared import Pt, Inches, RGBColor
        from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
        from docx.oxml.shared import OxmlElement, qn
        
        # Import the base directory setup from streamlined_analyzer
        import os
        base_dir = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
        reports_dir = os.path.join(base_dir, "Reports")
        os.makedirs(reports_dir, exist_ok=True)
        
        doc = Document()
        
        # Set up document styling
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Calibri'
        font.size = Pt(11)
        
        # Add custom bullet styles if they don't exist
        try:
            # Try to access existing bullet styles, create if needed
            list_bullet_style = doc.styles['List Bullet']
        except:
            # Create custom bullet style
            list_bullet_style = doc.styles.add_style('List Bullet', 1)  # 1 = paragraph style
            list_bullet_style.font.name = 'Calibri'
            list_bullet_style.font.size = Pt(11)
        
        try:
            list_bullet_2_style = doc.styles['List Bullet 2']
        except:
            list_bullet_2_style = doc.styles.add_style('List Bullet 2', 1)
            list_bullet_2_style.font.name = 'Calibri'
            list_bullet_2_style.font.size = Pt(10)
        
        try:
            list_bullet_3_style = doc.styles['List Bullet 3']
        except:
            list_bullet_3_style = doc.styles.add_style('List Bullet 3', 1)
            list_bullet_3_style.font.name = 'Calibri'
            list_bullet_3_style.font.size = Pt(10)
        
        # Title
        title_text = f"Kuwait Banking Headlines Report"
        title = doc.add_heading(title_text, 0)
        title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        # Format title
        title_run = title.runs[0]
        title_run.font.size = Pt(24)
        title_run.font.color.rgb = RGBColor(0, 51, 102)  # Dark blue
        
        # Subtitle with date and analysis type
        subtitle = doc.add_paragraph(f"{analysis_type.title()} Analysis • {datetime.now().strftime('%B %d, %Y')}")
        subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        subtitle_run = subtitle.runs[0]
        subtitle_run.font.size = Pt(14)
        subtitle_run.font.color.rgb = RGBColor(100, 100, 100)
        subtitle_run.italic = True
        
        # Add summary statistics
        doc.add_paragraph()  # spacing
        summary_para = doc.add_paragraph()
        summary_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        total_banks = len(bank_headlines)
        banks_with_coverage = sum(1 for bank_data in bank_headlines.values() 
                                 if bank_data['results']['statistics']['newspapersWithContent'] > 0)
        total_relevant_pages = sum(bank_data['results']['statistics']['totalRelevantPages'] 
                                  for bank_data in bank_headlines.values())
        
        summary_run = summary_para.add_run(f"Coverage Summary: {banks_with_coverage} of {total_banks} banks with media coverage • {total_relevant_pages} total relevant pages found")
        summary_run.font.size = Pt(12)
        summary_run.font.color.rgb = RGBColor(0, 100, 0)  # Green
        summary_run.bold = True
        
        # Add main content
        doc.add_paragraph()  # spacing
        doc.add_heading("Headlines", level=1)
        
        # Initialize analyzer for AI processing
        analyzer = CompanyAnalyzer()
        
        # Sort banks alphabetically
        sorted_banks = sorted(bank_headlines.items())
        
        for bank_name, bank_data in sorted_banks:
            results = bank_data['results']
            
            # Add bank heading with hyperlink styling
            bank_heading = doc.add_heading(bank_name, level=2)
            bank_heading_run = bank_heading.runs[0]
            bank_heading_run.font.color.rgb = RGBColor(0, 100, 200)  # Blue link color
            bank_heading_run.underline = True
            
            # Check if bank has any coverage
            if results['statistics']['newspapersWithContent'] == 0:
                no_coverage_para = doc.add_paragraph("• No media coverage found in analyzed newspapers")
                no_coverage_run = no_coverage_para.runs[0]
                no_coverage_run.font.italic = True
                no_coverage_run.font.color.rgb = RGBColor(150, 150, 150)
                doc.add_paragraph()  # spacing
                continue
            
            # Add coverage summary
            coverage_summary = doc.add_paragraph(f"• Found coverage in {results['statistics']['newspapersWithContent']} newspaper(s)")
            coverage_summary.style = 'List Bullet'
            coverage_run = coverage_summary.runs[0]
            coverage_run.font.color.rgb = RGBColor(0, 100, 0)  # Green
            coverage_run.bold = True
            
            # Extract actual headlines/news titles for this bank from each newspaper
            try:
                headlines_found = False
                
                for newspaper_result in results['newspaperResults']:
                    if newspaper_result['success'] and newspaper_result['pagesFound'] > 0:
                        newspaper_name = newspaper_result['newspaper']
                        
                        # Get the analysis content - try different field names
                        analysis_content = None
                        if 'analysis' in newspaper_result and newspaper_result['analysis']:
                            analysis_content = newspaper_result['analysis']
                        elif 'summary' in newspaper_result and newspaper_result['summary']:
                            analysis_content = newspaper_result['summary']
                        else:
                            # Look for any text content in the result
                            for key, value in newspaper_result.items():
                                if isinstance(value, str) and len(value) > 50 and bank_name.lower() in value.lower():
                                    analysis_content = value
                                    print(f"📝 Using content from field '{key}' for {bank_name} in {newspaper_name}")
                                    break
                        
                        if analysis_content and len(analysis_content.strip()) > 10:
                            print(f"📰 Processing {newspaper_name} for {bank_name} - Content length: {len(analysis_content)}")
                            
                            # Enhanced prompt for headline extraction
                            headlines_prompt = f"""
                            You are a news editor creating headlines for {bank_name} based on this newspaper content from {newspaper_name}.
                            
                            Content about {bank_name}:
                            {analysis_content[:800]}
                            
                            Create 1-3 concise, factual news headlines about {bank_name} based on this content.
                            
                            Guidelines:
                            - Write in newspaper headline style (clear, direct, actionable)
                            - Focus on what {bank_name} is doing, announcing, or involved in
                            - If the content mentions specific activities, events, or business developments, highlight those
                            - If content is general, create a relevant business headline
                            
                            Examples of good banking headlines:
                            • "NBK Wealth releases market analysis report"
                            • "Gulf Bank launches customer rewards program"
                            • "Commercial Bank expands digital services"
                            • "Kuwait Finance House reports quarterly performance"
                            • "Warba Bank sponsors community initiative"
                            
                            Return exactly 1-3 headlines in this format:
                            • [Headline 1]
                            • [Headline 2]
                            • [Headline 3]
                            
                            Be creative but factual. Make it sound like real news headlines about {bank_name}.
                            """
                            
                            try:
                                print(f"🤖 Calling AI for headlines: {bank_name} - {newspaper_name}")
                                newspaper_headlines = await analyzer.call_llm_with_retry(headlines_prompt)
                                print(f"🎯 AI Response for {bank_name}/{newspaper_name}: {newspaper_headlines[:200]}...")
                                
                                # Clean up the AI response and filter for actual headlines
                                headline_lines = []
                                for line in newspaper_headlines.split('\n'):
                                    line = line.strip()
                                    if line and (line.startswith('•') or line.startswith('-')):
                                        # Clean the line
                                        clean_line = line.replace('•', '').replace('-', '').strip()
                                        if len(clean_line) > 10:  # Must be substantial
                                            headline_lines.append(f"• {clean_line}")
                                
                                if headline_lines and len(headline_lines) > 0:
                                    print(f"✅ Found {len(headline_lines)} headlines for {bank_name} in {newspaper_name}")
                                    
                                    # Add newspaper subheading
                                    newspaper_subheading = doc.add_paragraph(f"From {newspaper_name}:")
                                    newspaper_subheading_run = newspaper_subheading.runs[0]
                                    newspaper_subheading_run.font.size = Pt(10)
                                    newspaper_subheading_run.font.color.rgb = RGBColor(0, 100, 200)  # Blue
                                    newspaper_subheading_run.italic = True
                                    newspaper_subheading_run.bold = True
                                    
                                    # Add each headline with proper indentation
                                    for headline in headline_lines[:3]:  # Max 3 headlines per newspaper
                                        headline_para = doc.add_paragraph(f"  {headline}")
                                        headline_run = headline_para.runs[0]
                                        headline_run.font.size = Pt(10)
                                        headline_run.font.color.rgb = RGBColor(60, 60, 60)  # Dark gray
                                    
                                    headlines_found = True
                                else:
                                    print(f"⚠️ AI returned no valid headlines for {bank_name} in {newspaper_name}")
                                    # Create a generic headline
                                    newspaper_subheading = doc.add_paragraph(f"From {newspaper_name}:")
                                    newspaper_subheading_run = newspaper_subheading.runs[0]
                                    newspaper_subheading_run.font.size = Pt(10)
                                    newspaper_subheading_run.font.color.rgb = RGBColor(0, 100, 200)
                                    newspaper_subheading_run.italic = True
                                    newspaper_subheading_run.bold = True
                                    
                                    generic_headline = f"• {bank_name} featured in business news coverage"
                                    headline_para = doc.add_paragraph(f"  {generic_headline}")
                                    headline_run = headline_para.runs[0]
                                    headline_run.font.size = Pt(10)
                                    headline_run.font.color.rgb = RGBColor(100, 100, 100)
                                    headline_run.italic = True
                                    
                                    headlines_found = True
                                    
                            except Exception as e:
                                print(f"❌ AI Error for {bank_name}/{newspaper_name}: {e}")
                                # Create fallback headline
                                newspaper_subheading = doc.add_paragraph(f"From {newspaper_name}:")
                                newspaper_subheading_run = newspaper_subheading.runs[0]
                                newspaper_subheading_run.font.size = Pt(10)
                                newspaper_subheading_run.font.color.rgb = RGBColor(0, 100, 200)
                                newspaper_subheading_run.italic = True
                                newspaper_subheading_run.bold = True
                                
                                fallback_headline = f"• {bank_name} coverage found in business section"
                                headline_para = doc.add_paragraph(f"  {fallback_headline}")
                                headline_run = headline_para.runs[0]
                                headline_run.font.size = Pt(10)
                                headline_run.font.color.rgb = RGBColor(120, 120, 120)
                                
                                headlines_found = True
                        else:
                            print(f"⚠️ No usable content for {bank_name} in {newspaper_name}")
                
                # If absolutely no headlines were found, add a fallback message
                if not headlines_found:
                    print(f"❌ No headlines found for {bank_name} - using final fallback")
                    fallback_para = doc.add_paragraph(f"• Content analysis in progress - coverage confirmed in {results['statistics']['newspapersWithContent']} newspaper(s)")
                    fallback_run = fallback_para.runs[0]
                    fallback_run.font.size = Pt(10)
                    fallback_run.font.italic = True
                    fallback_run.font.color.rgb = RGBColor(150, 150, 150)
                    
            except Exception as e:
                print(f"💥 Major error processing headlines for {bank_name}: {e}")
                import traceback
                traceback.print_exc()
                
                # Emergency fallback
                emergency_para = doc.add_paragraph(f"• Coverage processing error - found in {results['statistics']['newspapersWithContent']} newspaper(s)")
                emergency_run = emergency_para.runs[0]
                emergency_run.font.size = Pt(10)
                emergency_run.font.color.rgb = RGBColor(200, 100, 100)  # Reddish
            
            doc.add_paragraph()  # spacing between banks
        
        # Add footer section
        doc.add_page_break()
        doc.add_heading("Analysis Details", level=1)
        
        # Create summary table
        table = doc.add_table(rows=1, cols=4)
        table.style = 'Table Grid'
        
        # Headers
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Bank'
        hdr_cells[1].text = 'Newspapers with Coverage'
        hdr_cells[2].text = 'Total Pages Found'
        hdr_cells[3].text = 'Status'
        
        for cell in hdr_cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
        
        # Add data rows
        for bank_name, bank_data in sorted_banks:
            results = bank_data['results']
            row_cells = table.add_row().cells
            
            row_cells[0].text = bank_name
            row_cells[1].text = str(results['statistics']['newspapersWithContent'])
            row_cells[2].text = str(results['statistics']['totalRelevantPages'])
            
            if results['statistics']['newspapersWithContent'] > 0:
                row_cells[3].text = "✅ Coverage Found"
            else:
                row_cells[3].text = "➖ No Coverage"
        
        # Add generation info
        doc.add_paragraph()
        info_para = doc.add_paragraph(f"Report generated on {datetime.now().strftime('%B %d, %Y at %I:%M %p')}")
        info_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        info_run = info_para.runs[0]
        info_run.font.size = Pt(10)
        info_run.font.italic = True
        info_run.font.color.rgb = RGBColor(100, 100, 100)
        
        # Save report
        report_filename = f"Kuwait_Banking_Headlines_{analysis_type.title()}_{datetime.now().strftime('%Y-%m-%d')}.docx"
        report_path = os.path.join(reports_dir, report_filename)
        doc.save(report_path)
        
        print(f"✅ Headlines report saved: {report_path}")
        return report_path
        
    except Exception as e:
        print(f"❌ Error generating headlines report: {e}")
        traceback.print_exc()
@app.route('/api/analytics-report', methods=['POST'])
def generate_analytics_report():
    """Generate analytics report with charts for multiple completed analyses"""
    try:
        data = request.get_json()
        session_ids = data.get('sessionIds', [])
        
        if not session_ids:
            return jsonify({"error": "No session IDs provided"}), 400
        
        # Collect results from all sessions
        bank_analytics = {}
        analysis_type = "detailed"  # Default
        
        for session_id in session_ids:
            if session_id in analysis_sessions:
                session = analysis_sessions[session_id]
                if session.results and hasattr(session, 'company_name'):
                    bank_analytics[session.company_name] = {
                        'results': session.results,
                        'session': session
                    }
                    analysis_type = session.analysis_type
        
        if not bank_analytics:
            return jsonify({"error": "No valid completed analyses found"}), 400
        
        # Generate analytics report
        print(f"📊 Generating analytics report for {len(bank_analytics)} banks")
        
        # Run the async function in a new event loop
        def run_analytics_generation():
            loop = asyncio.new_event_loop()
            asyncio.set_event_loop(loop)
            try:
                report_path = loop.run_until_complete(
                    generate_analytics_report_with_charts(bank_analytics, analysis_type)
                )
                return report_path
            finally:
                loop.close()
        
        report_path = run_analytics_generation()
        
        if report_path and os.path.exists(report_path):
            filename = os.path.basename(report_path)
            
            return send_file(
                report_path,
                as_attachment=True,
                download_name=filename,
                mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            )
        else:
            return jsonify({"error": "Failed to generate analytics report"}), 500
            
    except Exception as e:
        print(f"❌ Analytics report error: {e}")
        traceback.print_exc()
        return jsonify({"error": f"Failed to generate analytics report: {str(e)}"}), 500

async def generate_analytics_report_with_charts(bank_analytics, analysis_type):
    """Generate analytics report with charts and insights"""
    try:
        from docx import Document
        from docx.shared import Pt, Inches, RGBColor
        from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
        import io
        import base64
        
        # Import chart generation libraries
        import matplotlib
        matplotlib.use('Agg')  # Use non-GUI backend
        import matplotlib.pyplot as plt
        import numpy as np
        
        # Import the base directory setup
        import os
        base_dir = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
        reports_dir = os.path.join(base_dir, "Reports")
        os.makedirs(reports_dir, exist_ok=True)
        
        doc = Document()
        
        # Set up document styling
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Calibri'
        font.size = Pt(11)
        
        # Title
        title_text = f"Kuwait Banking Analytics Report"
        title = doc.add_heading(title_text, 0)
        title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        # Format title
        title_run = title.runs[0]
        title_run.font.size = Pt(24)
        title_run.font.color.rgb = RGBColor(0, 51, 102)  # Dark blue
        
        # Subtitle
        subtitle = doc.add_paragraph(f"{analysis_type.title()} Analysis • {datetime.now().strftime('%B %d, %Y')}")
        subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        subtitle_run = subtitle.runs[0]
        subtitle_run.font.size = Pt(14)
        subtitle_run.font.color.rgb = RGBColor(100, 100, 100)
        subtitle_run.italic = True
        
        # Initialize analyzer for AI processing
        analyzer = CompanyAnalyzer()
        
        # Calculate analytics data
        press_releases_data = {}  # Unique coverage per bank
        total_hits_data = {}      # Total coverage per bank
        
        for bank_name, bank_data in bank_analytics.items():
            results = bank_data['results']
            
            # Calculate total hits (all coverage regardless of uniqueness)
            total_hits = results['statistics']['totalRelevantPages']
            total_hits_data[bank_name] = total_hits
            
            # Calculate unique press releases using AI analysis
            unique_count = await calculate_unique_press_releases(analyzer, bank_name, results)
            press_releases_data[bank_name] = unique_count
        
        # Generate Chart 1: Press Releases (Unique Coverage)
        chart1_path = generate_press_releases_chart(press_releases_data)
        
        # Generate Chart 2: Total Hits
        chart2_path = generate_total_hits_chart(total_hits_data)
        
        # Add summary statistics
        doc.add_paragraph()
        summary_para = doc.add_paragraph()
        summary_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        total_banks = len(bank_analytics)
        banks_with_coverage = sum(1 for count in total_hits_data.values() if count > 0)
        total_unique_releases = sum(press_releases_data.values())
        total_all_hits = sum(total_hits_data.values())
        
        summary_run = summary_para.add_run(
            f"Analysis Summary: {banks_with_coverage} of {total_banks} banks with coverage • "
            f"{total_unique_releases} unique press releases • {total_all_hits} total media hits"
        )
        summary_run.font.size = Pt(12)
        summary_run.font.color.rgb = RGBColor(0, 100, 0)
        summary_run.bold = True
        
        # Chart 1: Press Releases
        doc.add_paragraph()
        doc.add_heading("Chart 1: Total Number of Press Releases", level=1)
        
        chart1_desc = doc.add_paragraph(
            "This chart shows the number of unique press releases per bank. "
            "If the same news story appears in multiple newspapers, it counts as 1 press release."
        )
        chart1_desc.runs[0].font.italic = True
        chart1_desc.runs[0].font.size = Pt(10)
        
        if chart1_path and os.path.exists(chart1_path):
            doc.add_picture(chart1_path, width=Inches(6.0))
            last_paragraph = doc.paragraphs[-1]
            last_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        # Chart 2: Total Hits
        doc.add_page_break()
        doc.add_heading("Chart 2: Total Number of Hits", level=1)
        
        chart2_desc = doc.add_paragraph(
            "This chart shows the total number of media hits per bank across all newspapers. "
            "Each mention in each newspaper counts separately, regardless of uniqueness."
        )
        chart2_desc.runs[0].font.italic = True
        chart2_desc.runs[0].font.size = Pt(10)
        
        if chart2_path and os.path.exists(chart2_path):
            doc.add_picture(chart2_path, width=Inches(6.0))
            last_paragraph = doc.paragraphs[-1]
            last_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        # Data Table
        doc.add_paragraph()
        doc.add_heading("Detailed Analytics", level=1)
        
        # Create analytics table
        table = doc.add_table(rows=1, cols=4)
        table.style = 'Table Grid'
        
        # Headers
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Bank'
        hdr_cells[1].text = 'Unique Press Releases'
        hdr_cells[2].text = 'Total Media Hits'
        hdr_cells[3].text = 'Coverage Ratio'
        
        for cell in hdr_cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
        
        # Sort banks by total hits (descending)
        sorted_banks = sorted(bank_analytics.items(), key=lambda x: total_hits_data[x[0]], reverse=True)
        
        # Add data rows
        for bank_name, bank_data in sorted_banks:
            row_cells = table.add_row().cells
            
            unique_releases = press_releases_data[bank_name]
            total_hits = total_hits_data[bank_name]
            
            row_cells[0].text = bank_name
            row_cells[1].text = str(unique_releases)
            row_cells[2].text = str(total_hits)
            
            # Calculate coverage ratio (hits per unique release)
            if unique_releases > 0:
                ratio = round(total_hits / unique_releases, 1)
                row_cells[3].text = f"{ratio}x"
            else:
                row_cells[3].text = "0x"
        
        # Key Insights Section
        doc.add_paragraph()
        doc.add_heading("Key Insights", level=1)
        
        # Generate AI insights
        insights_prompt = f"""
        Analyze these Kuwait banking media analytics and provide 3-5 key insights:
        
        Data:
        {chr(10).join([f"{bank}: {press_releases_data[bank]} unique releases, {total_hits_data[bank]} total hits" 
                      for bank in sorted([bank for bank in bank_analytics.keys()])])}
        
        Provide insights about:
        1. Which banks have the most media presence
        2. Which banks have the highest coverage efficiency (hits per release)
        3. Any notable patterns or differences
        4. Market leadership indicators
        
        Format as bullet points:
        • [Insight 1]
        • [Insight 2]
        • [Insight 3]
        """
        
        try:
            insights_result = await analyzer.call_llm_with_retry(insights_prompt)
            insights_lines = [line.strip() for line in insights_result.split('\n') 
                            if line.strip() and line.strip().startswith('•')]
            
            for insight in insights_lines:
                insight_para = doc.add_paragraph(insight)
                insight_para.style = 'List Bullet'
        except Exception as e:
            print(f"⚠️ Could not generate AI insights: {e}")
            # Add basic insights
            doc.add_paragraph("• Media coverage analysis completed across all selected banks")
            doc.add_paragraph(f"• {banks_with_coverage} out of {total_banks} banks had media presence")
            doc.add_paragraph(f"• Total of {total_unique_releases} unique stories generated {total_all_hits} media hits")
        
        # Add generation info
        doc.add_paragraph()
        info_para = doc.add_paragraph(f"Report generated on {datetime.now().strftime('%B %d, %Y at %I:%M %p')}")
        info_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        info_run = info_para.runs[0]
        info_run.font.size = Pt(10)
        info_run.font.italic = True
        info_run.font.color.rgb = RGBColor(100, 100, 100)
        
        # Save report
        report_filename = f"Kuwait_Banking_Analytics_{analysis_type.title()}_{datetime.now().strftime('%Y-%m-%d')}.docx"
        report_path = os.path.join(reports_dir, report_filename)
        doc.save(report_path)
        
        # Clean up chart files
        try:
            if chart1_path and os.path.exists(chart1_path):
                os.remove(chart1_path)
            if chart2_path and os.path.exists(chart2_path):
                os.remove(chart2_path)
        except:
            pass
        
        print(f"✅ Analytics report saved: {report_path}")
        return report_path
        
    except Exception as e:
        print(f"❌ Error generating analytics report: {e}")
        traceback.print_exc()
        return None

async def calculate_unique_press_releases(analyzer, bank_name, results):
    """Calculate unique press releases using AI to identify duplicate stories"""
    try:
        # Collect all headlines/analysis content for this bank
        all_content = []
        for newspaper_result in results['newspaperResults']:
            if newspaper_result['success'] and newspaper_result['pagesFound'] > 0:
                if 'analysis' in newspaper_result and newspaper_result['analysis']:
                    all_content.append({
                        'newspaper': newspaper_result['newspaper'],
                        'content': newspaper_result['analysis'][:500]  # Limit content
                    })
        
        if not all_content:
            return 0
        
        # Use AI to identify unique stories
        uniqueness_prompt = f"""
        Analyze these news contents about {bank_name} and identify unique press releases/stories.
        
        Content from different newspapers:
        {chr(10).join([f"{item['newspaper']}: {item['content']}" for item in all_content])}
        
        Task: Count the number of UNIQUE press releases/news stories about {bank_name}.
        - If the same story appears in multiple newspapers, count it as 1
        - If different stories/announcements are found, count each separately
        
        Examples:
        - Same earnings report in 3 newspapers = 1 unique press release
        - Earnings report + new branch opening = 2 unique press releases
        
        Return ONLY a number representing unique press releases count.
        """
        
        unique_count_result = await analyzer.call_llm_with_retry(uniqueness_prompt)
        
        # Extract number from AI response
        import re
        numbers = re.findall(r'\d+', unique_count_result)
        if numbers:
            unique_count = int(numbers[0])
            # Sanity check - can't be more than total newspaper count
            max_possible = len(all_content)
            return min(unique_count, max_possible)
        else:
            # Fallback: assume each newspaper has unique content
            return len(all_content)
            
    except Exception as e:
        print(f"⚠️ Error calculating unique releases for {bank_name}: {e}")
        # Fallback: count newspapers with content
        return len([r for r in results['newspaperResults'] if r['success'] and r['pagesFound'] > 0])

def generate_press_releases_chart(press_releases_data):
    """Generate Press Releases chart using matplotlib"""
    try:
        import matplotlib.pyplot as plt
        import numpy as np
        
        # Sort data by value for better visualization
        sorted_data = dict(sorted(press_releases_data.items(), key=lambda x: x[1], reverse=True))
        
        banks = list(sorted_data.keys())
        values = list(sorted_data.values())
        
        # Create chart
        plt.figure(figsize=(12, 8))
        bars = plt.bar(range(len(banks)), values, color=[
            '#dc2626', '#1d4ed8', '#16a34a', '#ea580c', '#dc2626', 
            '#16a34a', '#0891b2', '#eab308', '#a855f7', '#6b7280'
        ][:len(banks)])
        
        plt.title('Total Number of Press Releases', fontsize=16, fontweight='bold', pad=20)
        plt.xlabel('Banks', fontsize=12)
        plt.ylabel('Total Press Releases', fontsize=12)
        
        # Set x-axis labels
        plt.xticks(range(len(banks)), [bank.replace(' ', '\n') for bank in banks], rotation=0, ha='center')
        
        # Add value labels on bars
        for i, bar in enumerate(bars):
            height = bar.get_height()
            if height > 0:
                plt.text(bar.get_x() + bar.get_width()/2., height + 0.05,
                        f'{int(height)}', ha='center', va='bottom', fontweight='bold')
        
        # Set y-axis to show integers only
        plt.gca().yaxis.set_major_locator(plt.MaxNLocator(integer=True))
        
        # Add grid
        plt.grid(axis='y', alpha=0.3)
        
        # Adjust layout
        plt.tight_layout()
        
        # Save chart
        chart_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'temp_press_releases_chart.png')
        plt.savefig(chart_path, dpi=300, bbox_inches='tight')
        plt.close()
        
        return chart_path
        
    except Exception as e:
        print(f"❌ Error generating press releases chart: {e}")
        return None

def generate_total_hits_chart(total_hits_data):
    """Generate Total Hits chart using matplotlib"""
    try:
        import matplotlib.pyplot as plt
        import numpy as np
        
        # Sort data by value for better visualization
        sorted_data = dict(sorted(total_hits_data.items(), key=lambda x: x[1], reverse=True))
        
        banks = list(sorted_data.keys())
        values = list(sorted_data.values())
        
        # Create chart
        plt.figure(figsize=(12, 8))
        bars = plt.bar(range(len(banks)), values, color=[
            '#dc2626', '#1d4ed8', '#16a34a', '#ea580c', '#dc2626', 
            '#16a34a', '#0891b2', '#eab308', '#a855f7', '#6b7280'
        ][:len(banks)])
        
        plt.title('Total Number of Hits', fontsize=16, fontweight='bold', pad=20)
        plt.xlabel('Banks', fontsize=12)
        plt.ylabel('Total Hits', fontsize=12)
        
        # Set x-axis labels
        plt.xticks(range(len(banks)), [bank.replace(' ', '\n') for bank in banks], rotation=0, ha='center')
        
        # Add value labels on bars
        for i, bar in enumerate(bars):
            height = bar.get_height()
            if height > 0:
                plt.text(bar.get_x() + bar.get_width()/2., height + 0.1,
                        f'{int(height)}', ha='center', va='bottom', fontweight='bold')
        
        # Set y-axis to show integers only
        plt.gca().yaxis.set_major_locator(plt.MaxNLocator(integer=True))
        
        # Add grid
        plt.grid(axis='y', alpha=0.3)
        
        # Adjust layout
        plt.tight_layout()
        
        # Save chart
        chart_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'temp_total_hits_chart.png')
        plt.savefig(chart_path, dpi=300, bbox_inches='tight')
        plt.close()
        
        return chart_path
        
    except Exception as e:
        print(f"❌ Error generating total hits chart: {e}")
        return None

def create_detailed_web_results(company_name, executive_summary, all_results, successful_count, content_count, total_pages, total_newspapers, report_path):
    """Create web results for detailed analysis"""
    return {
        "companyName": company_name,
        "analysisDate": datetime.now().strftime('%Y-%m-%d'),
        "analysisType": "detailed",
        "executiveSummary": executive_summary,
        "reportPath": report_path,
        "newspaperResults": [
            {
                "id": result["newspaper"].lower().replace(' ', '_').replace('-', '_'),
                "newspaper": result["newspaper"],
                "language": result.get("language", "N/A"),
                "success": result["success"],
                "pagesFound": len(result["relevant_pages"]) if result["success"] else 0,
                "status": (
                    "Content Found" if result["success"] and result["relevant_pages"] else
                    "No Content" if result["success"] else
                    "Failed"
                ),
                "analysis": result.get("analysis", result.get("error", "No analysis available")),
                "screenshots": [
                    {
                        "pageNumber": page_num + 1,
                        "description": f"Page {page_num + 1} - Contains {company_name} content"
                    }
                    for page_num in result["relevant_pages"]
                ] if result["success"] and result["relevant_pages"] else []
            }
            for result in all_results
        ],
        "statistics": {
            "totalNewspapers": total_newspapers,
            "successfulAnalyses": successful_count,
            "newspapersWithContent": content_count,
            "totalRelevantPages": total_pages
        }
    }

if __name__ == '__main__':
    print("🚀 Starting Company News Analyzer API Server...")
    print("=" * 60)
    print("📱 Web interface will connect to: http://localhost:5001")
    print("🔧 API endpoints available:")
    print("   • GET  /api/health - Health check")
    print("   • GET  /api/companies - List available companies")
    print("   • GET  /api/newspapers - List available newspapers")
    print("   • POST /api/analyze - Start analysis")
    print("   • POST /api/analysis/{id}/stop - Stop analysis")
    print("   • GET  /api/analysis/{id}/status - Check progress")
    print("   • GET  /api/analysis/{id}/results - Get results")
    print("   • GET  /api/analysis/{id}/download - Download report")
    print("   • POST /api/headlines-report - Generate consolidated headlines report")
    print("   • POST /api/analytics-report - Generate analytics report with charts")
    print("=" * 60)
    print("💡 Make sure your .env file contains OPENAI_API_KEY")
    print("💡 Press Ctrl+C to stop the server")
    print("")
    
    # Check if streamlined_analyzer can be imported
    try:
        from streamlined_analyzer import CompanyAnalyzer, NEWSPAPERS
        print(f"✅ Found {len(NEWSPAPERS)} newspapers in configuration")
        print(f"✅ Found {len(PREDEFINED_COMPANIES)} predefined companies")
        for company in PREDEFINED_COMPANIES:
            print(f"   • {company}")
    except ImportError:
        print("❌ Cannot import streamlined_analyzer.py - make sure it exists!")
        print("   Copy your existing streamlined code to 'streamlined_analyzer.py'")
        exit(1)
    
    app.run(debug=False, host="0.0.0.0", port=int(os.environ.get("PORT", 5000)))
    
    
    
    