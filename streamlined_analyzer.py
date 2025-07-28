#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Streamlined Company News Analysis - 9 Kuwaiti Newspapers
Analyzes company mentions across major Kuwaiti newspapers using AI
"""

import os
import nest_asyncio
from dotenv import load_dotenv
from datetime import datetime, timedelta
import asyncio
from playwright.async_api import async_playwright
import requests
import fitz  # PyMuPDF
from langchain_openai import ChatOpenAI
import re
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from PIL import Image
import unicodedata
import json
import random

# Setup
nest_asyncio.apply()
load_dotenv()

if not os.getenv('OPENAI_API_KEY'):
    raise ValueError('OPENAI_API_KEY is not set')

# Base directories - FIXED to use project directory
base_dir = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
downloads_base = os.path.join(base_dir, "Downloads")
reports_dir = os.path.join(base_dir, "Reports")
os.makedirs(downloads_base, exist_ok=True)
os.makedirs(reports_dir, exist_ok=True)

# Newspaper configurations
NEWSPAPERS = {
    "Arab Times": {
        "url": "https://www.arabtimesonline.com/news/category/e-paper/",
        "language": "English",
        "download_dir": os.path.join(downloads_base, "ArabTimes"),
        "method": "arabtimes_specific"
    },
    "Al-Seyassah": {
        "url": "https://alseyassah.com/%d8%b9%d8%af%d8%af-%d8%a7%d9%84%d9%8a%d9%88%d9%85/",
        "language": "Arabic",
        "download_dir": os.path.join(downloads_base, "AlSeyassah"),
        "method": "alseyassah_specific"
    },
    "Al-Rai": {
        "url": "https://www.alraimedia.com/newspaper/pdf",
        "language": "Arabic",
        "download_dir": os.path.join(downloads_base, "AlRaiMedia"),
        "method": "alrai_specific"
    },
    "Al-Anba": {
        "url": "https://www.alanba.com.kw/newspaper/",
        "language": "Arabic",
        "download_dir": os.path.join(downloads_base, "AlAnbaMedia"),
        "method": "alanba_specific"
    },
    "Kuwait Times": {
        "url": "https://kuwaittimes.com/archive-pdf/",
        "language": "English",
        "download_dir": os.path.join(downloads_base, "KuwaitTimes"),
        "method": "generic"
    },
    "Al-Qabas": {
        "url": "https://www.alqabas.com",
        "language": "Arabic",
        "download_dir": os.path.join(downloads_base, "AlQabasMedia"),
        "method": "generic"
    },
    "Al-Jarida": {
        "url": "https://www.aljarida.com/",
        "language": "Arabic",
        "download_dir": os.path.join(downloads_base, "AlJarida"),
        "method": "generic"
    },
    "Al-Nahar": {
        "url": "https://www.annaharkw.com/Home",
        "language": "Arabic",
        "download_dir": os.path.join(downloads_base, "AlNahar"),
        "method": "alnahar_specific"
    },
    "Al-Wasat": {
        "url": "https://www.alwasat.com.kw/",
        "language": "Arabic",
        "download_dir": os.path.join(downloads_base, "AlWasat"),
        "method": "alwasat_specific"
    }
}

# Create download directories
for config in NEWSPAPERS.values():
    os.makedirs(config["download_dir"], exist_ok=True)

class CompanyAnalyzer:
    def __init__(self):
        self.llm = ChatOpenAI(model="gpt-4o", temperature=0)
        self.today = datetime.now()
    
    async def call_llm_with_retry(self, prompt: str, max_retries: int = 3) -> str:
        """Call LLM with retry logic for rate limiting"""
        for attempt in range(max_retries):
            try:
                result = self.llm.invoke(prompt)
                return result.content
            except Exception as e:
                if "rate_limit_exceeded" in str(e) or "429" in str(e):
                    if attempt < max_retries - 1:
                        wait_time = (2 ** attempt) + random.uniform(0, 1)
                        print(f"⏱️ Rate limit hit, waiting {wait_time:.1f} seconds...")
                        await asyncio.sleep(wait_time)
                        continue
                raise e
    
    async def generate_company_variations(self, company_name: str) -> dict:
        """Generate company name variations using AI"""
        print(f"🤖 Generating variations for: {company_name}")
        
        is_arabic = any(ord(char) > 127 for char in company_name)
        language_hint = "Arabic" if is_arabic else "English"
        
        prompt = f"""
        Generate ALL possible variations of company name "{company_name}" (appears to be {language_hint}) 
        for newspaper analysis. BE VERY COMPREHENSIVE.

        For "{company_name}":
        - Include both full and partial matches
        - Arabic companies: include with/without "بنك", "شركة", "مجموعة" prefixes
        - Include abbreviations, acronyms, informal names
        - Include English/Arabic translations and transliterations
        - Include quoted versions: «name», "name"
        - Include corporate indicators: ك.س.ب., K.S.C.P., S.A.K.

        Example for "بنك الوطني":
        English: National Bank, NBK, National Bank of Kuwait, Kuwait National Bank
        Arabic: بنك الوطني, الوطني, البنك الوطني, مجموعة الوطني, «الوطني», "الوطني"

        Return ONLY a JSON object:
        {{
            "english_variations": ["variation1", "variation2", ...],
            "arabic_variations": ["variation1", "variation2", ...],
            "search_patterns": {{
                "english_high_confidence": ["\\\\bNational Bank\\\\b", "\\\\bNBK\\\\b"],
                "arabic_high_confidence": ["الوطني", "بنك.*الوطني", "الوطني.*بنك"]
            }},
            "exclusion_patterns": ["exclude1", "exclude2"]
        }}

        Make patterns flexible to catch real content, not overly strict.
        """
        
        try:
            response = await self.call_llm_with_retry(prompt)
            response_clean = response.strip()
            if response_clean.startswith('```json'):
                response_clean = response_clean[7:]
            if response_clean.endswith('```'):
                response_clean = response_clean[:-3]
            
            variations = json.loads(response_clean)
            
            # Add fallback patterns for Arabic companies
            if is_arabic and "الوطني" in company_name:
                variations["arabic_variations"].extend([
                    "الوطني", "بنك الوطني", "البنك الوطني", "مجموعة الوطني", 
                    "«الوطني»", '"الوطني"', "الوطني ك.س.ب."
                ])
                variations["search_patterns"]["arabic_high_confidence"].extend([
                    "الوطني", "بنك.*الوطني", "الوطني.*بنك"
                ])
            
            print(f"✅ Generated {len(variations['english_variations'])} English, {len(variations['arabic_variations'])} Arabic variations")
            return variations
            
        except Exception as e:
            print(f"❌ AI generation failed: {e}, using enhanced basic variations")
            return self.generate_enhanced_basic_variations(company_name)
    
    def generate_enhanced_basic_variations(self, company_name: str) -> dict:
        """Enhanced fallback with better Arabic support"""
        is_arabic = any(ord(char) > 127 for char in company_name)
        
        if is_arabic:
            # Enhanced Arabic variations
            arabic_variations = [company_name]
            english_variations = []
            
            if "الوطني" in company_name:
                arabic_variations.extend([
                    "الوطني", "بنك الوطني", "البنك الوطني", "مجموعة الوطني",
                    "«الوطني»", '"الوطني"', "الوطني ك.س.ب."
                ])
                english_variations.extend([
                    "National Bank", "NBK", "National Bank of Kuwait", 
                    "Kuwait National Bank", "Al Watani", "Alwatani"
                ])
            
            search_patterns = {
                "english_high_confidence": [r'\bNational Bank\b', r'\bNBK\b'],
                "arabic_high_confidence": ["الوطني", r'بنك.*الوطني', r'الوطني.*بنك']
            }
        else:
            # Enhanced English variations
            english_variations = [
                company_name, company_name.upper(), company_name.title(),
                f"The {company_name}", f"{company_name} Company", f"{company_name} Bank"
            ]
            arabic_variations = []
            
            search_patterns = {
                "english_high_confidence": [rf'\b{re.escape(company_name)}\b'],
                "arabic_high_confidence": [rf'\b{re.escape(company_name)}\b']
            }
        
        return {
            "english_variations": english_variations,
            "arabic_variations": arabic_variations,
            "search_patterns": search_patterns,
            "exclusion_patterns": []
        }
    
    def normalize_arabic_text(self, text: str) -> str:
        """Normalize Arabic text"""
        if not text:
            return text
        
        normalized = unicodedata.normalize('NFKC', text)
        replacements = {
            'ﺑ': 'ب', 'ﺒ': 'ب', 'ﻧ': 'ن', 'ﻨ': 'ن', 'ﻛ': 'ك', 'ﻜ': 'ك', 'ﻟ': 'ل', 'ﻠ': 'ل',
            '\u200c': '', '\u200d': '', '\u200e': '', '\u200f': '', '\xa0': ' '
        }
        
        for old, new in replacements.items():
            normalized = normalized.replace(old, new)
        
        arabic_diacritics = re.compile(r'[\u064B-\u0652\u0670\u0640]')
        normalized = arabic_diacritics.sub('', normalized)
        return normalized
    
    def get_potential_pages(self, page_texts: dict, variations: dict, language: str) -> list:
        """Find pages with potential company content"""
        potential_pages = []
        
        if language == "Arabic":
            patterns = variations['search_patterns']['arabic_high_confidence']
        else:
            patterns = variations['search_patterns']['english_high_confidence']
        
        exclusions = variations['exclusion_patterns']
        
        for page_num, text in page_texts.items():
            if language == "Arabic":
                text = self.normalize_arabic_text(text)
            
            score = sum(len(re.findall(pattern, text, re.UNICODE | re.IGNORECASE)) for pattern in patterns)
            exclusion_score = sum(len(re.findall(pattern, text, re.UNICODE | re.IGNORECASE)) for pattern in exclusions)
            
            if score > 0 and score > exclusion_score:
                potential_pages.append((page_num, score))
        
        potential_pages.sort(key=lambda x: x[1], reverse=True)
        return [page_num for page_num, score in potential_pages[:10]]
    
    async def download_pdf_generic(self, config: dict) -> str:
        """Generic PDF download method"""
        pdf_path = os.path.join(config["download_dir"], f"newspaper_{self.today.strftime('%Y-%m-%d')}.pdf")
        
        if os.path.exists(pdf_path):
            return pdf_path
        
        async with async_playwright() as p:
            browser = await p.chromium.launch(headless=True)
            page = await browser.new_page()
            
            try:
                await page.goto(config["url"], wait_until="networkidle", timeout=30000)
                await asyncio.sleep(2)
                
                content = await page.content()
                pdf_matches = re.findall(r'https?://[^\s"]+\.pdf[^\s"]*', content, re.IGNORECASE)
                
                for pdf_url in pdf_matches:
                    try:
                        response = requests.get(pdf_url, stream=True, timeout=30)
                        if response.status_code == 200 and 'pdf' in response.headers.get('content-type', '').lower():
                            with open(pdf_path, 'wb') as f:
                                for chunk in response.iter_content(chunk_size=8192):
                                    f.write(chunk)
                            
                            # Verify PDF
                            doc = fitz.open(pdf_path)
                            if len(doc) > 0:
                                doc.close()
                                await browser.close()
                                return pdf_path
                            doc.close()
                    except:
                        continue
                
                await browser.close()
            except:
                await browser.close()
        
        return None
    
    async def download_arabtimes_specific(self, config: dict) -> str:
        """Arab Times specific download - working method"""
        pdf_path = os.path.join(config["download_dir"], f"ArabTimes_{self.today.strftime('%Y-%m-%d')}.pdf")
        
        if os.path.exists(pdf_path):
            return pdf_path
        
        print("📱 Downloading Arab Times newspaper...")
        
        async with async_playwright() as p:
            browser = await p.chromium.launch(headless=True)
            context = await browser.new_context(viewport={"width": 1920, "height": 1080})
            page = await context.new_page()
            
            try:
                # Navigate to e-paper page
                await page.goto(config["url"], wait_until="networkidle", timeout=30000)
                await asyncio.sleep(2)
                
                # Find most recent edition using improved selector
                recent_edition = await page.evaluate("""() => {
                    const articles = document.querySelectorAll('article, .post, .entry, .item');
                    
                    for (const article of articles) {
                        const rect = article.getBoundingClientRect();
                        if (rect.width > 100 && rect.height > 100) {
                            const link = article.querySelector('a');
                            if (link && link.href) {
                                return {
                                    x: rect.x + rect.width/2,
                                    y: rect.y + rect.height/2,
                                    href: link.href
                                };
                            }
                        }
                    }
                    
                    // Fallback: look for e-paper links
                    const links = Array.from(document.querySelectorAll('a'));
                    for (const link of links) {
                        if (link.href && (link.href.includes('e-paper') || 
                                         link.textContent.toLowerCase().includes('e-paper') ||
                                         link.href.includes('2025'))) {
                            const rect = link.getBoundingClientRect();
                            if (rect.width > 0 && rect.height > 0) {
                                return {
                                    x: rect.x + rect.width/2,
                                    y: rect.y + rect.height/2,
                                    href: link.href
                                };
                            }
                        }
                    }
                    return null;
                }""")
                
                if recent_edition:
                    await page.mouse.click(recent_edition['x'], recent_edition['y'])
                    await page.wait_for_load_state("networkidle", timeout=15000)
                    await asyncio.sleep(2)
                    
                    # Find newspaper thumbnail to click
                    thumbnail_element = await page.evaluate("""() => {
                        const images = Array.from(document.querySelectorAll('img'));
                        const candidates = [];
                        
                        for (const img of images) {
                            const rect = img.getBoundingClientRect();
                            if (rect.width > 200 && rect.height > 200) {
                                candidates.push({
                                    x: rect.x + rect.width/2,
                                    y: rect.y + rect.height/2,
                                    area: rect.width * rect.height
                                });
                            }
                        }
                        
                        candidates.sort((a, b) => b.area - a.area);
                        return candidates.length > 0 ? candidates[0] : null;
                    }""")
                    
                    if thumbnail_element:
                        await page.mouse.click(thumbnail_element['x'], thumbnail_element['y'])
                        await asyncio.sleep(3)
                        
                        # Search for PDF URL in page content
                        page_content = await page.content()
                        pdf_matches = re.findall(r'https?://[^\s"]+\.pdf[^\s"]*', page_content, re.IGNORECASE)
                        
                        for pdf_url in pdf_matches:
                            if 'arabtimesonline.com' in pdf_url:
                                try:
                                    response = requests.get(pdf_url, stream=True, timeout=30)
                                    if response.status_code == 200 and 'pdf' in response.headers.get('content-type', '').lower():
                                        with open(pdf_path, 'wb') as f:
                                            for chunk in response.iter_content(chunk_size=8192):
                                                f.write(chunk)
                                        
                                        # Verify PDF
                                        try:
                                            doc = fitz.open(pdf_path)
                                            page_count = len(doc)
                                            doc.close()
                                            if page_count > 0:
                                                await browser.close()
                                                return pdf_path
                                        except:
                                            os.remove(pdf_path)
                                except:
                                    continue
                
                await browser.close()
            except Exception as e:
                print(f"❌ Error during download: {e}")
                await browser.close()
        
        print("❌ Failed to download PDF")
        return None
    
    async def download_alseyassah_specific(self, config: dict) -> str:
        """Al-Seyassah specific download - working method"""
        pdf_path = os.path.join(config["download_dir"], f"AlSeyassah_{self.today.strftime('%Y-%m-%d')}.pdf")
        
        if os.path.exists(pdf_path):
            return pdf_path
        
        print("📱 Downloading Al-Seyassah newspaper...")
        
        async with async_playwright() as p:
            browser = await p.chromium.launch(headless=True)
            context = await browser.new_context(viewport={"width": 1920, "height": 1080})
            page = await context.new_page()
            
            try:
                # Navigate to archive page
                await page.goto(config["url"], wait_until="networkidle", timeout=30000)
                await asyncio.sleep(2)
                
                # Find most recent edition using improved method
                recent_edition = await page.evaluate("""() => {
                    const candidates = [];
                    const clickableElements = Array.from(document.querySelectorAll('a, div[onclick], span[onclick], img[onclick], .clickable, [role="button"]'));
                    
                    for (const el of clickableElements) {
                        const rect = el.getBoundingClientRect();
                        if (rect.width > 50 && rect.height > 50 && rect.top < window.innerHeight) {
                            const text = el.textContent || el.alt || el.title || '';
                            const href = el.href || el.getAttribute('onclick') || '';
                            
                            if (text.match(/\\d{4}-\\d{2}-\\d{2}/) || text.match(/\\d{2}\\/\\d{2}\\/\\d{4}/) || 
                                href.includes('article') || href.includes('2025') || 
                                text.includes('2025') || el.tagName === 'IMG') {
                                
                                candidates.push({
                                    x: rect.x + rect.width/2,
                                    y: rect.y + rect.height/2,
                                    right: rect.right,
                                    href: href
                                });
                            }
                        }
                    }
                    
                    candidates.sort((a, b) => b.right - a.right);
                    return candidates.length > 0 ? candidates[0] : null;
                }""")
                
                if recent_edition:
                    await page.mouse.click(recent_edition['x'], recent_edition['y'])
                    await page.wait_for_load_state("networkidle", timeout=15000)
                    await asyncio.sleep(2)
                    
                    # Get page content and search for PDF URLs with multiple patterns
                    page_content = await page.content()
                    
                    pdf_patterns = [
                        r'https?://[^"\s]+\.pdf[^"\s]*',
                        r'https?://[^"\s]*pdf[^"\s]*\.pdf', 
                        r'/[^"\s]*\.pdf[^"\s]*',  # Relative URLs
                        r'["\']([^"\']*\.pdf[^"\']*)["\']',
                    ]
                    
                    found_pdf_urls = []
                    for pattern in pdf_patterns:
                        matches = re.findall(pattern, page_content, re.IGNORECASE)
                        found_pdf_urls.extend(matches)
                    
                    unique_urls = list(set(found_pdf_urls))
                    
                    for url in unique_urls:
                        # Convert relative URLs to absolute
                        if url.startswith('/'):
                            url = 'https://alseyassah.com' + url
                        
                        if 'alseyassah.com' in url:
                            try:
                                response = requests.get(url, stream=True, timeout=30)
                                if response.status_code == 200:
                                    content_type = response.headers.get('content-type', '').lower()
                                    if 'pdf' in content_type or url.endswith('.pdf'):
                                        with open(pdf_path, 'wb') as f:
                                            for chunk in response.iter_content(chunk_size=8192):
                                                f.write(chunk)
                                        
                                        # Verify PDF
                                        try:
                                            doc = fitz.open(pdf_path)
                                            page_count = len(doc)
                                            doc.close()
                                            if page_count > 0:
                                                await browser.close()
                                                return pdf_path
                                        except:
                                            os.remove(pdf_path)
                                            continue
                            except:
                                continue
                else:
                    print("❌ No recent edition found")
            
            except Exception as e:
                print(f"❌ Error during download: {e}")
            
            finally:
                await browser.close()
        
        print("❌ Failed to download PDF")
        return None
    
    async def download_alrai_specific(self, config: dict) -> str:
        """Al-Rai specific download - working method"""
        pdf_path = os.path.join(config["download_dir"], f"AlRai_{self.today.strftime('%Y-%m-%d')}.pdf")
        
        if os.path.exists(pdf_path):
            return pdf_path
        
        print("📱 Downloading Al-Rai newspaper...")
        
        async with async_playwright() as p:
            browser = await p.chromium.launch(headless=False)
            context = await browser.new_context(viewport={"width": 1920, "height": 1080})
            page = await context.new_page()
            
            try:
                # Navigate to PDF viewer page
                await page.goto(config["url"], wait_until="networkidle", timeout=15000)
                await asyncio.sleep(3)
                
                # Click "إضغط هنا" button to open PDF viewer
                click_here_elements = await page.locator("text=إضغط هنا").all()
                
                if click_here_elements:
                    await click_here_elements[0].click()
                    await page.wait_for_load_state("networkidle", timeout=10000)
                    await asyncio.sleep(8)
                    
                    # Check if redirected to direct PDF URL
                    current_url = page.url
                    
                    if current_url.endswith('.pdf') or '.pdf?' in current_url:
                        try:
                            response = requests.get(current_url, stream=True, timeout=30)
                            if response.status_code == 200:
                                with open(pdf_path, 'wb') as f:
                                    for chunk in response.iter_content(chunk_size=8192):
                                        f.write(chunk)
                                
                                # Verify PDF
                                if os.path.exists(pdf_path) and os.path.getsize(pdf_path) > 10000:
                                    try:
                                        doc = fitz.open(pdf_path)
                                        page_count = len(doc)
                                        doc.close()
                                        if page_count > 0:
                                            await browser.close()
                                            return pdf_path
                                    except:
                                        os.remove(pdf_path)
                        except:
                            pass
                
                await browser.close()
            except Exception as e:
                await browser.close()
        
        return None
    
    async def download_alanba_specific(self, config: dict) -> str:
        """Al-Anba specific download"""
        pdf_path = os.path.join(config["download_dir"], f"AlAnba_{self.today.strftime('%Y-%m-%d')}.pdf")
        
        if os.path.exists(pdf_path):
            return pdf_path
        
        # Try direct download with multiple date formats
        dates_to_try = [self.today, self.today - timedelta(days=1)]
        
        for date_attempt in dates_to_try:
            date_str = date_attempt.strftime('%d-%m-%Y')
            year = date_attempt.strftime('%Y')
            month = date_attempt.strftime('%m')
            
            pdf_url = f"https://pdf.alanba.com.kw/pdf/{year}/{month}/{date_str}/{date_str}.pdf"
            
            try:
                response = requests.get(pdf_url, stream=True, timeout=30)
                if response.status_code == 200:
                    with open(pdf_path, 'wb') as f:
                        for chunk in response.iter_content(chunk_size=8192):
                            f.write(chunk)
                    
                    # Verify
                    doc = fitz.open(pdf_path)
                    if len(doc) > 0:
                        doc.close()
                        return pdf_path
                    doc.close()
            except:
                continue
        
        return None
    
    async def download_alnahar_specific(self, config: dict) -> str:
        """Al-Nahar specific download - working method"""
        pdf_path = os.path.join(config["download_dir"], f"AlNahar_{self.today.strftime('%Y-%m-%d')}.pdf")
        
        if os.path.exists(pdf_path):
            return pdf_path
        
        print("📱 Downloading Al-Nahar newspaper...")
        
        async with async_playwright() as p:
            browser = await p.chromium.launch(headless=False)
            context = await browser.new_context(viewport={"width": 1920, "height": 1080})
            page = await context.new_page()
            
            try:
                # Navigate to the main page
                main_url = "https://www.annaharkw.com/Home"
                await page.goto(main_url, wait_until="domcontentloaded", timeout=45000)
                await asyncio.sleep(5)
                
                # Look for clickable elements in the top area
                clickable_candidates = await page.evaluate(f"""() => {{
                    const candidates = [];
                    const clickableElements = Array.from(document.querySelectorAll('a, div[onclick], span[onclick], img[onclick], .clickable, [role="button"], button'));
                    
                    for (const el of clickableElements) {{
                        const rect = el.getBoundingClientRect();
                        if (rect.width > 10 && rect.height > 10 && rect.top < 200) {{
                            const text = el.textContent || el.alt || el.title || '';
                            const href = el.href || el.getAttribute('onclick') || '';
                            
                            const isPdfRelated = (
                                text.toLowerCase().includes('pdf') ||
                                href.toLowerCase().includes('pdf') ||
                                href.includes('2025') ||
                                href.includes('{self.today.strftime('%Y%m%d')}')
                            );
                            
                            if (isPdfRelated || (rect.width < 100 && rect.height < 100 && rect.top < 100)) {{
                                candidates.push({{
                                    x: rect.x + rect.width/2,
                                    y: rect.y + rect.height/2,
                                    text: text.slice(0, 100),
                                    href: href,
                                    isPdfRelated: isPdfRelated
                                }});
                            }}
                        }}
                    }}
                    
                    candidates.sort((a, b) => {{
                        if (a.isPdfRelated && !b.isPdfRelated) return -1;
                        if (!a.isPdfRelated && b.isPdfRelated) return 1;
                        return a.x - b.x;
                    }});
                    
                    return candidates;
                }}""")
                
                # Try clicking on the most promising candidates
                pdf_downloaded = False
                
                for candidate in clickable_candidates[:5]:
                    if pdf_downloaded:
                        break
                        
                    try:
                        await page.mouse.click(candidate['x'], candidate['y'])
                        await asyncio.sleep(3)
                        
                        # Check for new pages/tabs
                        pages = context.pages
                        if len(pages) > 1:
                            new_page = pages[-1]
                            new_url = new_page.url
                            
                            # Check if it's a PDF URL and not an old one
                            if (new_url.endswith('.pdf') or 'pdf' in new_url.lower()) and not any(old_year in new_url for old_year in ['2023', '2022', '2024', '2021']):
                                pdf_downloaded = await self.try_download_pdf(new_url, pdf_path, context, new_page)
                                if pdf_downloaded:
                                    break
                            
                            await new_page.close()
                        
                        # Check if current page URL changed to PDF
                        current_url = page.url
                        if (current_url.endswith('.pdf') or 'pdf' in current_url.lower()) and not any(old_year in current_url for old_year in ['2023', '2022', '2024', '2021']):
                            pdf_downloaded = await self.try_download_pdf(current_url, pdf_path, context, page)
                            if pdf_downloaded:
                                break
                        
                    except Exception:
                        continue
                
                await browser.close()
                
                if pdf_downloaded:
                    return pdf_path
                
            except Exception as e:
                await browser.close()
        
        return None
    
    async def try_download_pdf(self, pdf_url: str, pdf_path: str, context, page=None) -> bool:
        """Try multiple methods to download PDF"""
        # Method 1: Direct browser download with session
        try:
            cookies = await context.cookies()
            session = requests.Session()
            
            for cookie in cookies:
                session.cookies.set(cookie['name'], cookie['value'], domain=cookie.get('domain', ''))
            
            headers = {
                'User-Agent': 'Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36',
                'Referer': 'https://www.annaharkw.com/',
                'Accept': 'application/pdf,application/octet-stream,*/*',
                'Accept-Language': 'en-US,en;q=0.9,ar;q=0.8',
                'Accept-Encoding': 'gzip, deflate, br',
                'Connection': 'keep-alive',
                'Upgrade-Insecure-Requests': '1'
            }
            
            response = session.get(pdf_url, headers=headers, stream=True, timeout=30)
            
            if response.status_code == 200:
                content_type = response.headers.get('content-type', '').lower()
                if 'pdf' in content_type or pdf_url.endswith('.pdf'):
                    with open(pdf_path, 'wb') as f:
                        for chunk in response.iter_content(chunk_size=8192):
                            f.write(chunk)
                    
                    if await self.validate_pdf(pdf_path):
                        return True
        except Exception:
            pass
        
        # Method 2: JavaScript fetch (if page is available)
        if page:
            try:
                pdf_content = await page.evaluate(f"""
                    async () => {{
                        try {{
                            const response = await fetch('{pdf_url}');
                            if (response.ok) {{
                                const arrayBuffer = await response.arrayBuffer();
                                const bytes = new Uint8Array(arrayBuffer);
                                return Array.from(bytes);
                            }}
                        }} catch (e) {{
                            console.error('Fetch failed:', e);
                        }}
                        return null;
                    }}
                """)
                
                if pdf_content:
                    with open(pdf_path, 'wb') as f:
                        f.write(bytes(pdf_content))
                    if await self.validate_pdf(pdf_path):
                        return True
            except Exception:
                pass
        
        return False
    
    async def validate_pdf(self, pdf_path: str) -> bool:
        """Validate downloaded PDF"""
        try:
            if os.path.exists(pdf_path):
                file_size = os.path.getsize(pdf_path)
                if file_size > 10000:  # At least 10KB
                    doc = fitz.open(pdf_path)
                    pages = len(doc)
                    doc.close()
                    if pages > 0:
                        return True
        except Exception:
            pass
        return False
    
    async def download_alwasat_specific(self, config: dict) -> str:
        """Al-Wasat specific download"""
        pdf_path = os.path.join(config["download_dir"], f"AlWasat_{self.today.strftime('%Y-%m-%d')}.pdf")
        
        if os.path.exists(pdf_path):
            return pdf_path
        
        async with async_playwright() as p:
            browser = await p.chromium.launch(headless=True)
            page = await browser.new_page()
            
            try:
                await page.goto(config["url"], wait_until="domcontentloaded", timeout=30000)
                await asyncio.sleep(3)
                
                # Look for PDF links
                pdf_links = await page.evaluate("""() => {
                    const links = Array.from(document.querySelectorAll('a, [href]'));
                    const pdfLinks = [];
                    
                    for (const element of links) {
                        const href = element.href || '';
                        if (href.includes('.pdf')) {
                            const rect = element.getBoundingClientRect();
                            if (rect.width > 0 && rect.height > 0) {
                                pdfLinks.push({href: href});
                            }
                        }
                    }
                    return pdfLinks;
                }""")
                
                for link in pdf_links:
                    try:
                        response = requests.get(link['href'], stream=True, timeout=30)
                        if response.status_code == 200 and 'pdf' in response.headers.get('content-type', '').lower():
                            with open(pdf_path, 'wb') as f:
                                for chunk in response.iter_content(chunk_size=8192):
                                    f.write(chunk)
                            
                            # Verify
                            doc = fitz.open(pdf_path)
                            if len(doc) > 0:
                                doc.close()
                                await browser.close()
                                return pdf_path
                            doc.close()
                    except:
                        continue
                
                await browser.close()
            except:
                await browser.close()
        
        return None
    
    async def download_pdf(self, newspaper_name: str, config: dict) -> str:
        """Download PDF using appropriate method"""
        method = config.get("method", "generic")
        
        if method == "arabtimes_specific":
            return await self.download_arabtimes_specific(config)
        elif method == "alseyassah_specific":
            return await self.download_alseyassah_specific(config)
        elif method == "alrai_specific":
            return await self.download_alrai_specific(config)
        elif method == "alanba_specific":
            return await self.download_alanba_specific(config)
        elif method == "alnahar_specific":
            return await self.download_alnahar_specific(config)
        elif method == "alwasat_specific":
            return await self.download_alwasat_specific(config)
        else:
            return await self.download_pdf_generic(config)
    
    async def analyze_newspaper(self, newspaper_name: str, config: dict, company_name: str, variations: dict) -> dict:
        """Analyze newspaper for company content"""
        print(f"🔍 Analyzing {newspaper_name} for {company_name}...")
        
        pdf_path = await self.download_pdf(newspaper_name, config)
        if not pdf_path:
            return {
                "newspaper": newspaper_name,
                "success": False,
                "error": "Failed to download PDF",
                "relevant_pages": [],
                "analysis": ""
            }
        
        try:
            doc = fitz.open(pdf_path)
            page_texts = {}
            page_screenshots = {}
            
            # Extract text and create screenshots
            for page_num in range(len(doc)):
                page = doc[page_num]
                text = page.get_text()
                
                if config["language"] == "Arabic":
                    text = self.normalize_arabic_text(text)
                
                page_texts[page_num] = text
                
                # Create screenshot
                pix = page.get_pixmap(matrix=fitz.Matrix(2, 2))
                screenshot_path = os.path.join(config["download_dir"], f"page_{page_num+1}.png")
                pix.save(screenshot_path)
                page_screenshots[page_num] = screenshot_path
            
            doc.close()
            
            # Find potential pages
            potential_pages = self.get_potential_pages(page_texts, variations, config["language"])
            
            if not potential_pages:
                # Fallback to business section
                total_pages = len(page_texts)
                start = max(5, total_pages // 3)
                end = min(total_pages, start + 5)
                potential_pages = list(range(start, end))
            
            # Analyze pages with AI
            relevant_pages = []
            page_analyses = {}
            
            for page_num in potential_pages:
                page_prompt = f"""
                Analyze this {config['language']} newspaper text for mentions of {company_name} or related entities.
                
                COMPANY VARIATIONS TO LOOK FOR:
                English: {', '.join(variations['english_variations'][:8])}
                Arabic: {', '.join(variations['arabic_variations'][:8])}
                
                Text: {page_texts[page_num][:3000]}
                
                IMPORTANT: Look for ANY mention of the company, even if it's a partial match or related entity.
                For "بنك الوطني" also look for just "الوطني" when referring to the bank.
                For banking companies, look for financial activities, bond issuances, banking services, etc.
                
                BE FLEXIBLE - if you find content related to the company (even with slight variations), include it.
                
                If you find ANY relevant content, respond with:
                ## SUMMARY
                [What you found about the company/bank]
                
                ## SENTIMENT
                [Positive/Negative/Neutral]
                
                FOUND: {company_name} content detected on this page.
                
                If NO relevant content found, respond only with: "NO CONTENT"
                """
                
                try:
                    analysis = await self.call_llm_with_retry(page_prompt)
                    
                    if "FOUND:" in analysis and "NO CONTENT" not in analysis:
                        relevant_pages.append(page_num)
                        page_analyses[page_num] = analysis
                    elif "NO CONTENT" not in analysis:
                        # Fallback check for any mention in the analysis
                        company_terms = [company_name] + variations['english_variations'][:3] + variations['arabic_variations'][:3]
                        if any(term.lower() in analysis.lower() for term in company_terms if term):
                            relevant_pages.append(page_num)
                            page_analyses[page_num] = analysis
                    
                except Exception as e:
                    print(f"❌ Error analyzing page {page_num+1}: {e}")
                    continue
            
            # Create combined analysis - ONLY include pages with actual content
            if relevant_pages:
                combined_analysis = f"Found {company_name} content on {len(relevant_pages)} page(s) with actual mentions:\n\n"
                for page_num in sorted(relevant_pages):
                    combined_analysis += f"--- Page {page_num+1} ---\n{page_analyses[page_num]}\n\n"
            else:
                combined_analysis = f"No {company_name} content found in {newspaper_name}."
            
            # Clean up PDF
            try:
                os.remove(pdf_path)
            except:
                pass
            
            return {
                "newspaper": newspaper_name,
                "success": True,
                "relevant_pages": relevant_pages,
                "analysis": combined_analysis,
                "page_screenshots": page_screenshots,
                "language": config["language"]
            }
            
        except Exception as e:
            print(f"❌ Error analyzing {newspaper_name}: {e}")
            return {
                "newspaper": newspaper_name,
                "success": False,
                "error": str(e),
                "relevant_pages": [],
                "analysis": ""
            }
    
    def resize_image(self, image_path: str, max_width: int = 500) -> str:
        """Resize image for document"""
        try:
            with Image.open(image_path) as img:
                width, height = img.size
                if width > max_width:
                    ratio = max_width / width
                    new_height = int(height * ratio)
                    resized_img = img.resize((max_width, new_height), Image.Resampling.LANCZOS)
                    resized_path = image_path.replace('.png', '_resized.png')
                    resized_img.save(resized_path)
                    return resized_path
            return image_path
        except:
            return image_path
    
    async def create_report(self, all_results: list, company_name: str, analysis_type: str = "detailed") -> str:
        """Create comprehensive Word report"""
        doc = Document()
        
        # Title
        title_text = f"{company_name} {analysis_type.title()} Media Analysis Report"
        title = doc.add_heading(title_text, 0)
        title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        subtitle = doc.add_paragraph(f"Analysis Date: {self.today.strftime('%B %d, %Y')}")
        subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        # Executive Summary
        doc.add_heading("📊 Executive Summary", level=1)
        
        successful_analyses = [r for r in all_results if r["success"] and r["relevant_pages"]]
        total_newspapers = len(all_results)
        newspapers_with_content = len(successful_analyses)
        
        summary_para = doc.add_paragraph()
        summary_para.add_run("Analysis Overview: ").bold = True
        summary_para.add_run(f"Analyzed {total_newspapers} newspapers. Found {company_name} content in {newspapers_with_content} publications.")
        
        if successful_analyses:
            # Create executive summary using AI
            try:
                all_content = [f"{r['newspaper']}: {r['analysis']}" for r in successful_analyses]
                
                if analysis_type == "short":
                    # Short analysis - ONE paragraph with bullet points only
                    exec_prompt = f"""
                    Create a VERY BRIEF summary of {company_name}'s media coverage in ONE paragraph with bullet points.
                    Keep it simple and factual, no analysis or sentiment.
                    
                    Content from newspapers:
                    {chr(10).join(all_content)}
                    
                    Provide ONLY a single paragraph with bullet points of the main news items about {company_name}.
                    Format: "Key developments include: • [item 1] • [item 2] • [item 3]"
                    Focus only on facts, no sentiment or business implications.
                    """
                    
                    executive_summary = await self.call_llm_with_retry(exec_prompt)
                    doc.add_paragraph(executive_summary.strip())
                    
                else:
                    # Detailed analysis - full summary with sections
                    exec_prompt = f"""
                    Create executive summary for {company_name}'s media coverage:
                    {chr(10).join(all_content)}
                    
                    Provide:
                    ## KEY DEVELOPMENTS
                    [Main stories about {company_name}]
                    
                    ## SENTIMENT
                    [Overall sentiment]
                    
                    ## INSIGHTS
                    [Business implications]
                    """
                    
                    executive_summary = await self.call_llm_with_retry(exec_prompt)
                    
                    sections = executive_summary.split("##")
                    for section in sections:
                        if not section.strip():
                            continue
                        
                        lines = section.strip().split("\n", 1)
                        if len(lines) >= 2:
                            section_name = lines[0].strip()
                            section_content = lines[1].strip()
                            
                            if section_name:
                                doc.add_heading(section_name, level=2)
                            doc.add_paragraph(section_content)
                        
            except Exception as e:
                print(f"Error creating executive summary: {e}")
                if analysis_type == "short":
                    doc.add_paragraph(f"Found {company_name} content in {newspapers_with_content} out of {total_newspapers} newspapers analyzed.")
                else:
                    doc.add_paragraph("Executive summary could not be generated.")
        else:
            if analysis_type == "short":
                doc.add_paragraph(f"No {company_name} content found in the {total_newspapers} newspapers analyzed.")
            else:
                doc.add_paragraph(f"No {company_name} content found in analyzed newspapers.")
        
        # Coverage Summary Table
        doc.add_heading("📋 Coverage Summary", level=1)
        
        table = doc.add_table(rows=1, cols=4)
        table.style = 'Table Grid'
        
        # Headers
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Newspaper'
        hdr_cells[1].text = 'Language'
        hdr_cells[2].text = 'Status'
        hdr_cells[3].text = 'Pages Found'
        
        for cell in hdr_cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
        
        # Add data
        for result in all_results:
            row_cells = table.add_row().cells
            row_cells[0].text = result["newspaper"]
            row_cells[1].text = result.get("language", "N/A")
            
            if result["success"]:
                if result["relevant_pages"]:
                    row_cells[2].text = "✅ Content Found"
                    row_cells[3].text = f"{len(result['relevant_pages'])} pages"
                else:
                    row_cells[2].text = "✅ No Content"
                    row_cells[3].text = "0 pages"
            else:
                row_cells[2].text = "❌ Failed"
                row_cells[3].text = "N/A"
        
        # Individual newspaper sections - ONLY show newspapers with actual content
        for result in all_results:
            if not result["success"] or not result["relevant_pages"]:
                continue
            
            doc.add_page_break()
            doc.add_heading(f"📰 {result['newspaper']}", level=1)
            
            # Add newspaper info
            info_para = doc.add_paragraph()
            info_para.add_run(f"Language: {result.get('language', 'N/A')} | ").italic = True
            info_para.add_run(f"Pages with {company_name} content: {len(result['relevant_pages'])}").italic = True
            info_para.runs[-1].font.color.rgb = RGBColor(0, 128, 0)  # Green
            
            # Analysis - adapt content based on analysis type
            doc.add_heading("Analysis", level=2)
            
            if analysis_type == "short":
                # For short analysis, provide ONLY 2-3 sentences summary
                try:
                    analysis_text = result["analysis"]
                    short_summary_prompt = f"""
                    Summarize this newspaper analysis in EXACTLY 2-3 sentences with key facts only:
                    
                    {analysis_text}
                    
                    Provide only the main facts about {company_name} found in this newspaper.
                    No sentiment, no business implications, just the basic facts.
                    Keep it to 2-3 sentences maximum.
                    """
                    
                    short_summary = await self.call_llm_with_retry(short_summary_prompt)
                    doc.add_paragraph(short_summary.strip())
                    
                except Exception as e:
                    print(f"Error creating short summary: {e}")
                    # Fallback: just take first 200 characters
                    analysis_text = result["analysis"]
                    if len(analysis_text) > 200:
                        condensed_analysis = analysis_text[:200] + "..."
                    else:
                        condensed_analysis = analysis_text
                    doc.add_paragraph(condensed_analysis)
            else:
                # For detailed analysis, include full analysis
                doc.add_paragraph(result["analysis"])
            
            # Screenshots of relevant pages - INCLUDED IN BOTH SHORT AND DETAILED
            if result.get("page_screenshots") and result["relevant_pages"]:
                doc.add_heading(f"Pages with {company_name} Content", level=2)
                
                # Only show screenshots for pages that actually have content
                for page_num in sorted(result["relevant_pages"]):
                    if page_num in result["page_screenshots"]:
                        screenshot_path = result["page_screenshots"][page_num]
                        
                        if os.path.exists(screenshot_path):
                            page_para = doc.add_paragraph()
                            page_para.add_run(f"Page {page_num+1} - ").bold = True
                            page_para.add_run(f"{result['newspaper']} ({self.today.strftime('%Y-%m-%d')})")
                            page_para.add_run(f" - Contains {company_name} content").italic = True
                            
                            resized_path = self.resize_image(screenshot_path)
                            
                            try:
                                doc.add_picture(resized_path, width=Inches(5.0))
                                doc.add_paragraph()
                            except Exception as e:
                                print(f"Error adding image: {e}")
                                doc.add_paragraph(f"[Screenshot unavailable: {screenshot_path}]")
        
        # Save report
        company_safe_name = re.sub(r'[^\w\s-]', '', company_name).strip().replace(' ', '_')
        report_filename = f"{company_safe_name}_{analysis_type.title()}_Analysis_{self.today.strftime('%Y-%m-%d')}.docx"
        report_path = os.path.join(reports_dir, report_filename)
        doc.save(report_path)
        
        return report_path
    
    async def cleanup_files(self, all_results: list):
        """Clean up temporary files"""
        print("🧹 Cleaning up temporary files...")
        
        for result in all_results:
            if result["success"] and "page_screenshots" in result:
                for screenshot_path in result["page_screenshots"].values():
                    try:
                        if os.path.exists(screenshot_path):
                            os.remove(screenshot_path)
                        resized_path = screenshot_path.replace('.png', '_resized.png')
                        if os.path.exists(resized_path):
                            os.remove(resized_path)
                    except:
                        pass

def get_company_input() -> str:
    """Get company name from user"""
    print("\n" + "="*50)
    print("🏢 COMPANY NEWS ANALYZER")
    print("="*50)
    print("Analyzes company mentions across 9 Kuwaiti newspapers")
    
    while True:
        company_name = input("\n📝 Enter company name: ").strip()
        
        if not company_name or len(company_name) < 2:
            print("❌ Please enter a valid company name (at least 2 characters)")
            continue
        
        print(f"\n✅ Company: {company_name}")
        print("🤖 AI will generate variations and analyze content")
        
        confirm = input("🤔 Proceed? (y/n): ").strip().lower()
        if confirm in ['y', 'yes', '']:
            return company_name
        elif confirm in ['n', 'no']:
            continue
        else:
            print("Please enter 'y' or 'n'")

async def main():
    """Main execution function"""
    print("🚀 Starting Company News Analysis")
    
    # Get company name
    company_name = get_company_input()
    
    # Initialize analyzer
    analyzer = CompanyAnalyzer()
    
    # Generate variations
    print(f"\n🤖 Generating variations for: {company_name}")
    variations = await analyzer.generate_company_variations(company_name)
    
    print(f"\n📅 Analysis Date: {analyzer.today.strftime('%Y-%m-%d')}")
    print(f"🏢 Company: {company_name}")
    print(f"📰 Newspapers: {len(NEWSPAPERS)}")
    print("="*50)
    
    # Analyze each newspaper
    all_results = []
    
    for newspaper_name, config in NEWSPAPERS.items():
        try:
            result = await analyzer.analyze_newspaper(newspaper_name, config, company_name, variations)
            all_results.append(result)
            
            if result["success"]:
                if result["relevant_pages"]:
                    print(f"✅ {newspaper_name}: Found content on {len(result['relevant_pages'])} pages")
                else:
                    print(f"✅ {newspaper_name}: No content found")
            else:
                print(f"❌ {newspaper_name}: Failed - {result.get('error', 'Unknown error')}")
                
        except Exception as e:
            print(f"💥 {newspaper_name}: Fatal error - {e}")
            all_results.append({
                "newspaper": newspaper_name,
                "success": False,
                "error": str(e),
                "relevant_pages": [],
                "analysis": ""
            })
    
    # Generate report
    print(f"\n📝 Generating report for {company_name}...")
    try:
        report_path = await analyzer.create_report(all_results, company_name)
        print(f"✅ Report saved: {report_path}")
        
        # Summary statistics
        successful_count = sum(1 for r in all_results if r["success"])
        content_count = sum(1 for r in all_results if r["success"] and r["relevant_pages"])
        total_pages = sum(len(r["relevant_pages"]) for r in all_results if r["success"])
        
        print(f"\n📊 SUMMARY:")
        print(f"   📰 Newspapers analyzed: {len(all_results)}")
        print(f"   ✅ Successful: {successful_count}")
        print(f"   📄 With {company_name} content: {content_count}")
        print(f"   📋 Total relevant pages: {total_pages}")
        
        if content_count > 0:
            print(f"\n🎯 Coverage found in:")
            for result in all_results:
                if result["success"] and result["relevant_pages"]:
                    print(f"   • {result['newspaper']}: {len(result['relevant_pages'])} pages")
        
    except Exception as e:
        print(f"❌ Error generating report: {e}")
    
    # Cleanup
    await analyzer.cleanup_files(all_results)
    
    print(f"\n🏁 Analysis completed for {company_name}!")

if __name__ == "__main__":
    try:
        asyncio.run(main())
    except KeyboardInterrupt:
        print("\n⏹️ Analysis interrupted")
    except Exception as e:
        print(f"\n💥 Error: {e}")
    finally:
        print("🎬 Done.");